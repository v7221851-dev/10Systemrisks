# -*- coding: utf-8 -*-
"""Выгрузка описания алгоритма (средневзвешенная) в Word."""
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_para_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def main():
    doc = Document()
    doc.add_heading("Как работает алгоритм: средневзвешенные формулы", 0)

    add_para(doc, "В расчёте три уровня: балл по фактору (1–5), балл по группе рисков (средневзвешенный), итоговый индекс (тоже средневзвешенный по группам).")

    add_heading(doc, "1. Балл по одному фактору (1–5)", level=1)
    add_para(doc, "Каждый показатель сначала переводится в балл по шкале 1–5 (5 = норма, 1 = плохо).")
    add_para_bullet(doc, "Числовой показатель (range): используется функция calculate_risk — значение в пределах нормы [norm_min, norm_max] даёт 5.0; за пределами нормы балл линейно снижается до 1.0.")
    add_para_bullet(doc, "Категориальный (select): выбор пользователя маппится в балл 1–5 (например, «Норма» → 5, «Высокий риск» → 1).")
    add_para(doc, "Пример (числовой): артериальное давление 120/80, норма 110–130 → значение в норме → балл 5.0. Давление 150 → за нормой → балл, например, 2.2.")
    add_para(doc, "В расчётах дальше везде используются уже эти баллы 1–5.")

    add_heading(doc, "2. Балл группы рисков — средневзвешенный по факторам", level=1)
    add_para(doc, "По каждой группе (Neuro, Cardio, Hormone и т.д.) считается один балл — средневзвешенное баллов факторов с учётом весов из базы (Weight_Coefficient). В каждой группе веса факторов в сумме дают 1.0.")
    add_para(doc, "Формула в коде (по сути классическая средневзвешенная):")
    add_para(doc, "raw_score = Σ (балл_фактора × вес_фактора)")
    add_para(doc, "total_w   = Σ вес_фактора   (только по факторам, для которых есть ответ)")
    add_para(doc, "group_score = raw_score / total_w")
    add_para(doc, "То есть: group_score = (Σ (балл_i × w_i)) / (Σ w_i)")
    add_para(doc, "Пример для группы Cardio (упрощённо, 4 фактора, веса в сумме 1.0):", bold=True)
    add_para(doc, "Фактор: Давление — Вес (w_i): 0.35 — Балл пользователя: 5.0")
    add_para(doc, "Фактор: Холестерин — Вес: 0.30 — Балл: 3.0")
    add_para(doc, "Фактор: ЧСС покоя — Вес: 0.20 — Балл: 4.0")
    add_para(doc, "Фактор: С-реактивный белок — Вес: 0.15 — Балл: 5.0")
    add_para(doc, "raw_score = 5×0.35 + 3×0.30 + 4×0.20 + 5×0.15 = 1.75 + 0.90 + 0.80 + 0.75 = 4.20")
    add_para(doc, "total_w = 0.35 + 0.30 + 0.20 + 0.15 = 1.00")
    add_para(doc, "group_score(Cardio) = 4.20 / 1.00 = 4.20", bold=True)

    add_heading(doc, "3. Итоговый индекс — средневзвешенное по группам (с учётом «худшей»)", level=1)
    add_para(doc, "Итоговый балл считается не как простое среднее по группам, а как взвешенная комбинация среднего и минимума по группам (чтобы сильнее учитывать самую слабую систему).")
    add_para(doc, "В коде:")
    add_para_bullet(doc, "Если групп с данными меньше порога (например, 5): final_score = 0.7 × avg_score + 0.3 × min_score")
    add_para_bullet(doc, "Если групп достаточно: final_score = 0.6 × avg_score + 0.4 × min_score")
    add_para(doc, "avg_score — среднее арифметическое баллов по всем учтённым группам; min_score — минимальный балл среди этих групп (самая проблемная система).")
    add_para(doc, "Пример. Допустим по группам: Neuro 4.5, Cardio 4.2, Hormone 3.8, Metabolic 4.0, Immune 4.1 (достаточно групп).")
    add_para(doc, "avg_score = (4.5 + 4.2 + 3.8 + 4.0 + 4.1) / 5 = 4.12")
    add_para(doc, "min_score = 3.8 (Hormone)")
    add_para(doc, "final_score = 0.6 × 4.12 + 0.4 × 3.8 = 2.472 + 1.52 = 3.99 (округлённо 4.0)")
    add_para(doc, "Итоговый балл тоже в шкале 1–5.")

    add_heading(doc, "4. Перевод в проценты для интерфейса", level=1)
    add_para(doc, "Итоговый балл 1–5 переводится в проценты линейно: percent = final_score × 20 (1 → 20%, 5 → 100%). В коде это делает score_to_percent.")
    add_para(doc, "Пример: final_score = 3.99 → percent = 3.99 × 20 ≈ 80% (зелёная зона).")

    add_heading(doc, "Кратко по формулам", level=1)
    add_para(doc, "Уровень «Фактор»: балл 1–5 из значения/выбора — calculate_risk или маппинг опций.")
    add_para(doc, "Уровень «Группа рисков»: балл группы (1–5) — (Σ балл_i × w_i) / Σ w_i")
    add_para(doc, "Уровень «Итог»: итоговый балл (1–5) — 0.6×avg(group_scores) + 0.4×min(group_scores) (при достаточном числе групп).")
    add_para(doc, "Интерфейс: процент — final_score × 20")
    add_para(doc, "Итого: средневзвешенная используется явно на уровне группы рисков (по факторам), а на уровне итога — взвешенная комбинация среднего и минимума по группам.")

    out_path = "Алгоритм_средневзвешенная_формула.docx"
    doc.save(out_path)
    print(f"Сохранено: {out_path}")
    return out_path

if __name__ == "__main__":
    main()
